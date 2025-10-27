# Arquivo: core/views.py (COMPLETO E CORRIGIDO - Caminhos de Template de Auth)

# --- Importações Nativas e do Django ---
import os, io, locale, json, csv
from datetime import date, datetime
from django.shortcuts import render, redirect, get_object_or_404
from django.db.models import Count
from django.http import HttpResponse, JsonResponse, FileResponse
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.template.loader import render_to_string
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm

# --- Importações de Bibliotecas Externas ---
from weasyprint import HTML
from docxtpl import DocxTemplate, InlineImage
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
from decimal import Decimal, InvalidOperation
import google.generativeai as genai

# --- Importações do Projeto ---
from .models import (
    Processo, Orgao, Secretaria, Fornecedor, 
    Responsavel, HistoricoProcesso, Documento
)
from .forms import (
    ProcessoForm, OrgaoForm, FornecedorForm, 
    ResponsavelForm, ETPForm
)

# ==============================================================================
# FUNÇÃO AUXILIAR
# ==============================================================================
def _formata_moeda(valor):
    if valor in (None, ""): 
        return "Não informado"
    try: 
        v = Decimal(str(valor))
    except (InvalidOperation, ValueError, TypeError): 
        return "Não informado"
    reais, centavos = int(v), int((v - int(v)) * 100)
    reais_fmt = f"{reais:,d}".replace(",", ".")
    return f"R$ {reais_fmt},{centavos:02d}"

# ==============================================================================
# VIEW DO DASHBOARD
# ==============================================================================
@login_required
def dashboard(request):
    total_processos = Processo.objects.count()
    total_orgaos = Orgao.objects.count()
    total_fornecedores = Fornecedor.objects.count()
    processos_ativos = Processo.objects.filter(status__in=['FASE_INTERNA', 'PUBLICADO', 'AGUARDANDO_PROPOSTAS', 'EM_ANALISE']).count()
    processos_homologados = Processo.objects.filter(status='HOMOLOGADO').count()
    status_data = Processo.objects.values('status').annotate(count=Count('status')).order_by('status')
    status_choices = dict(Processo.STATUS_CHOICES)
    chart_labels = [status_choices.get(item['status'], item['status']) for item in status_data]
    chart_data = [item['count'] for item in status_data]
    atividades_recentes = HistoricoProcesso.objects.select_related('processo', 'processo__orgao_responsavel').order_by('-data_conclusao')[:5]
    context = {
        'total_processos': total_processos, 
        'total_orgaos': total_orgaos, 
        'total_fornecedores': total_fornecedores, 
        'processos_ativos': processos_ativos, 
        'processos_homologados': processos_homologados, 
        'chart_labels': chart_labels, 
        'chart_data': chart_data, 
        'atividades_recentes': atividades_recentes
    }
    return render(request, 'core/dashboard.html', context)

# ==============================================================================
# VIEWS DE AUTENTICAÇÃO (CAMINHOS CORRIGIDOS)
# ==============================================================================
def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('dashboard')
    else:
        form = AuthenticationForm()
    # CORRIGIDO o caminho do template:
    return render(request, 'core/login.html', {'form': form})

def logout_view(request):
    logout(request)
    return redirect('login')

def register_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('dashboard')
    else:
        form = UserCreationForm()
    # CORRIGIDO o caminho do template:
    return render(request, 'core/register.html', {'form': form})


# ==============================================================================
# VIEWS DE PROCESSOS
# ==============================================================================
@login_required
def processo_list(request):
    processos = Processo.objects.select_related('orgao_responsavel', 'secretaria_responsavel').order_by('-data_abertura')
    return render(request, 'core/processo_list.html', {'processos': processos})

@login_required
def processo_detail(request, pk):
    processo = get_object_or_404(Processo.objects.select_related('orgao_responsavel', 'secretaria_responsavel', 'responsavel_demanda'), pk=pk)
    historico_qs = HistoricoProcesso.objects.filter(processo=processo)
    etapas_concluidas_dict = {item.etapa: item.data_conclusao for item in historico_qs}
    todas_etapas = HistoricoProcesso.ETAPAS_CHOICES
    fase_interna_etapas = [{'key': key, 'name': name, 'date': etapas_concluidas_dict.get(key)} for key, name in todas_etapas[:12]]
    fase_externa_etapas = [{'key': key, 'name': name, 'date': etapas_concluidas_dict.get(key)} for key, name in todas_etapas[12:]]
    context = {
        'processo': processo, 
        'fase_interna_etapas': fase_interna_etapas, 
        'fase_externa_etapas': fase_externa_etapas
    }
    return render(request, 'core/processo_detail.html', context)

@login_required
def processo_documentos(request, pk):
    processo = get_object_or_404(Processo, pk=pk)
    return render(request, 'core/processo_documentos.html', {'processo': processo})

@login_required
def processo_create(request):
    if request.method == 'POST':
        form = ProcessoForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('processo_list')
    else:
        form = ProcessoForm()
    return render(request, 'core/processo_form.html', {'form': form, 'titulo': 'Adicionar Novo Processo'})

@login_required
def processo_update(request, pk):
    processo = get_object_or_404(Processo, pk=pk)
    if request.method == 'POST':
        form = ProcessoForm(request.POST, instance=processo)
        if form.is_valid():
            form.save()
            return redirect('processo_list')
    else:
        form = ProcessoForm(instance=processo)
    return render(request, 'core/processo_form.html', {'form': form, 'titulo': 'Editar Processo'})

@login_required
def processo_delete(request, pk):
    processo = get_object_or_404(Processo, pk=pk)
    if request.method == 'POST':
        processo.delete()
        return redirect('processo_list')
    return render(request, 'core/processo_confirm_delete.html', {'processo': processo})

@login_required
def processo_etp_form(request, pk):
    processo = get_object_or_404(Processo, pk=pk)
    if request.method == 'POST':
        form = ETPForm(request.POST, instance=processo)
        if form.is_valid():
            form.save()
            return redirect('processo_detail', pk=processo.pk)
    else:
        form = ETPForm(instance=processo)
    context = {
        'form': form, 
        'processo': processo, 
        'titulo': 'Estudo Técnico Preliminar (ETP)'
    }
    return render(request, 'core/etp_form.html', context)

# --- VIEWS DE GERAÇÃO DE DOCUMENTOS ---
@login_required
def documentos_list(request):
    """Lista todos os documentos gerados"""
    documentos = Documento.objects.all().select_related('processo')
    processos = Processo.objects.all()
    
    context = {
        'documentos': documentos,
        'processos': processos,
    }
    return render(request, 'core/documentos_list.html', context)

@login_required
def gerar_dfd(request, processo_id):
    """Gera o Documento de Formalização da Demanda"""
    processo = get_object_or_404(Processo, pk=processo_id)
    
    # Criar documento Word
    doc = DocxDocument()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título
    titulo = doc.add_heading('DOCUMENTO DE FORMALIZAÇÃO DA DEMANDA (DFD)', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações do Órgão
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Órgão: ').bold = True
    p.add_run(processo.orgao_responsavel.nome)
    
    if processo.secretaria_responsavel:
        p = doc.add_paragraph()
        p.add_run('Unidade Requisitante: ').bold = True
        p.add_run(processo.secretaria_responsavel.nome)
    
    p = doc.add_paragraph()
    p.add_run('Processo nº: ').bold = True
    p.add_run(processo.numero_processo)
    
    p = doc.add_paragraph()
    p.add_run('Data: ').bold = True
    p.add_run(datetime.now().strftime('%d/%m/%Y'))
    
    # 1. Descrição da Necessidade
    doc.add_heading('1. DESCRIÇÃO DA NECESSIDADE', 1)
    doc.add_paragraph(processo.justificativa or 'A ser preenchido.')
    
    # 2. Objeto da Contratação
    doc.add_heading('2. OBJETO DA CONTRATAÇÃO', 1)
    doc.add_paragraph(processo.descricao_detalhada_objeto or processo.objeto)
    
    # 3. Modalidade Sugerida
    doc.add_heading('3. MODALIDADE SUGERIDA', 1)
    doc.add_paragraph(f'{processo.get_modalidade_display()}')
    
    # 4. Valor Estimado
    doc.add_heading('4. VALOR ESTIMADO', 1)
    if processo.valor_estimado:
        doc.add_paragraph(f'R$ {processo.valor_estimado:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'))
    else:
        doc.add_paragraph('A ser definido após pesquisa de preços.')
    
    # 5. Justificativa
    doc.add_heading('5. JUSTIFICATIVA DA CONTRATAÇÃO', 1)
    doc.add_paragraph(processo.etp_justificativa_contratacao or 'A contratação se faz necessária para atender as demandas do órgão.')
    
    # 6. Responsável pela Demanda
    doc.add_heading('6. RESPONSÁVEL PELA DEMANDA', 1)
    if processo.responsavel_demanda:
        doc.add_paragraph(f'Nome: {processo.responsavel_demanda.nome}')
        doc.add_paragraph(f'Cargo: {processo.responsavel_demanda.cargo}')
        doc.add_paragraph(f'Matrícula: {processo.responsavel_demanda.matricula}')
    else:
        doc.add_paragraph('A ser definido.')
    
    # Assinatura
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('Assinatura do Responsável')
    
    # Salvar documento
    filename = f'DFD_{processo.numero_processo.replace("/", "-")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    filepath = os.path.join(settings.MEDIA_ROOT, 'documentos', filename)
    
    # Criar diretório se não existir
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc.save(filepath)
    
    # Salvar registro no banco
    documento = Documento.objects.create(
        processo=processo,
        tipo='DFD',
        arquivo=f'documentos/{filename}',
        gerado_por=request.user.username
    )
    
    messages.success(request, f'DFD gerado com sucesso!')
    
    # Retornar arquivo para download
    return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=filename)

@login_required
def gerar_etp(request, processo_id):
    """Gera o Estudo Técnico Preliminar"""
    processo = get_object_or_404(Processo, pk=processo_id)
    
    # Criar documento Word
    doc = DocxDocument()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título
    titulo = doc.add_heading('ESTUDO TÉCNICO PRELIMINAR (ETP)', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitulo = doc.add_heading('Conforme Art. 18, §1º da Lei 14.133/2021', 2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações do Órgão
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Órgão: ').bold = True
    p.add_run(processo.orgao_responsavel.nome)
    
    p = doc.add_paragraph()
    p.add_run('Processo nº: ').bold = True
    p.add_run(processo.numero_processo)
    
    p = doc.add_paragraph()
    p.add_run('Data: ').bold = True
    p.add_run(datetime.now().strftime('%d/%m/%Y'))
    
    # 1. Descrição da Necessidade
    doc.add_heading('1. DESCRIÇÃO DA NECESSIDADE', 1)
    doc.add_paragraph(processo.justificativa or 'A ser preenchido.')
    
    # 2. Previsão no PCA
    doc.add_heading('2. PREVISÃO NO PLANO DE CONTRATAÇÕES ANUAL (PCA)', 1)
    doc.add_paragraph(processo.etp_pca_texto)
    
    # 3. Descrição dos Requisitos da Contratação
    doc.add_heading('3. DESCRIÇÃO DOS REQUISITOS DA CONTRATAÇÃO', 1)
    doc.add_heading('3.1. Objeto da Contratação', 2)
    doc.add_paragraph(processo.descricao_detalhada_objeto or processo.objeto)
    
    if processo.etp_requisitos_tecnicos_detalhe:
        doc.add_heading('3.2. Requisitos Técnicos', 2)
        doc.add_paragraph(processo.etp_requisitos_tecnicos_detalhe)
    
    # 4. Levantamento de Mercado
    doc.add_heading('4. LEVANTAMENTO DE MERCADO', 1)
    doc.add_paragraph(processo.etp_texto_levantamento_mercado)
    
    # 5. Estimativa das Quantidades
    doc.add_heading('5. ESTIMATIVA DAS QUANTIDADES', 1)
    doc.add_paragraph(processo.etp_texto_estimativa_quantidades)
    
    # 6. Estimativa de Preços
    doc.add_heading('6. ESTIMATIVA DE PREÇOS', 1)
    if processo.valor_estimado:
        doc.add_paragraph(f'Valor Estimado Total: R$ {processo.valor_estimado:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'))
    if processo.etp_estimativa_metodologia:
        doc.add_paragraph(f'Metodologia: {processo.etp_estimativa_metodologia}')
    
    # 7. Descrição da Solução
    doc.add_heading('7. DESCRIÇÃO DA SOLUÇÃO COMO UM TODO', 1)
    doc.add_paragraph(processo.etp_descricao_solucao_texto)
    
    # 8. Justificativa para Parcelamento
    doc.add_heading('8. JUSTIFICATIVA PARA PARCELAMENTO OU NÃO DA SOLUÇÃO', 1)
    doc.add_paragraph(processo.etp_justificativa_parcelamento_texto)
    
    # 9. Contratações Correlatas
    doc.add_heading('9. CONTRATAÇÕES CORRELATAS E/OU INTERDEPENDENTES', 1)
    doc.add_paragraph(processo.etp_contratacoes_correlatas_texto)
    
    # 10. Alinhamento Estratégico
    doc.add_heading('10. ALINHAMENTO AO PLANEJAMENTO', 1)
    doc.add_paragraph(processo.etp_alinhamento_estrategico_texto or 'A contratação está alinhada com os objetivos estratégicos do órgão.')
    
    # 11. Resultados Pretendidos
    doc.add_heading('11. DEMONSTRATIVO DOS RESULTADOS PRETENDIDOS', 1)
    doc.add_paragraph(processo.etp_resultados_pretendidos_texto)
    
    # 12. Providências
    doc.add_heading('12. PROVIDÊNCIAS A SEREM ADOTADAS', 1)
    doc.add_paragraph(processo.etp_providencias_texto)
    
    # 13. Impactos Ambientais
    doc.add_heading('13. POSSÍVEIS IMPACTOS AMBIENTAIS', 1)
    doc.add_paragraph(processo.etp_impactos_ambientais_texto)
    
    # 14. Modalidade e Critério de Julgamento
    doc.add_heading('14. MODALIDADE E CRITÉRIO DE JULGAMENTO', 1)
    doc.add_paragraph(f'Modalidade: {processo.get_modalidade_display()}')
    if processo.etp_criterio_julgamento:
        doc.add_paragraph(f'Critério de Julgamento: {processo.get_etp_criterio_julgamento_display()}')
    if processo.etp_justificativa_modalidade_criterio:
        doc.add_paragraph(f'Justificativa: {processo.etp_justificativa_modalidade_criterio}')
    
    # 15. Prazo e Vigência
    doc.add_heading('15. PRAZO DE EXECUÇÃO E VIGÊNCIA CONTRATUAL', 1)
    if processo.vigencia_meses:
        doc.add_paragraph(f'Vigência: {processo.vigencia_meses} meses')
    if processo.etp_justificativa_prazo:
        doc.add_paragraph(f'Justificativa: {processo.etp_justificativa_prazo}')
    
    # 16. Dotação Orçamentária
    doc.add_heading('16. DOTAÇÃO ORÇAMENTÁRIA', 1)
    if processo.etp_dotacao_programa_trabalho:
        doc.add_paragraph(f'Programa de Trabalho: {processo.etp_dotacao_programa_trabalho}')
    if processo.etp_dotacao_natureza_despesa:
        doc.add_paragraph(f'Natureza da Despesa: {processo.etp_dotacao_natureza_despesa}')
    if processo.etp_dotacao_fonte_recursos:
        doc.add_paragraph(f'Fonte de Recursos: {processo.etp_dotacao_fonte_recursos}')
    
    # Assinatura
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    if processo.etp_responsavel_elaboracao:
        doc.add_paragraph(f'{processo.etp_responsavel_elaboracao.nome}')
        doc.add_paragraph(f'{processo.etp_responsavel_elaboracao.cargo}')
    else:
        doc.add_paragraph('Responsável pela Elaboração')
    
    # Salvar documento
    filename = f'ETP_{processo.numero_processo.replace("/", "-")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    filepath = os.path.join(settings.MEDIA_ROOT, 'documentos', filename)
    
    # Criar diretório se não existir
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc.save(filepath)
    
    # Salvar registro no banco
    documento = Documento.objects.create(
        processo=processo,
        tipo='ETP',
        arquivo=f'documentos/{filename}',
        gerado_por=request.user.username
    )
    
    messages.success(request, f'ETP gerado com sucesso!')
    
    # Retornar arquivo para download
    return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=filename)

@login_required
def gerar_tr(request, processo_id):
    """Gera o Termo de Referência"""
    processo = get_object_or_404(Processo, pk=processo_id)
    
    # Criar documento Word
    doc = DocxDocument()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título
    titulo = doc.add_heading('TERMO DE REFERÊNCIA', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações do Órgão
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Órgão: ').bold = True
    p.add_run(processo.orgao_responsavel.nome)
    
    p = doc.add_paragraph()
    p.add_run('Processo nº: ').bold = True
    p.add_run(processo.numero_processo)
    
    p = doc.add_paragraph()
    p.add_run('Data: ').bold = True
    p.add_run(datetime.now().strftime('%d/%m/%Y'))
    
    # 1. Objeto
    doc.add_heading('1. DO OBJETO', 1)
    doc.add_paragraph(processo.descricao_detalhada_objeto or processo.objeto)
    
    # 2. Justificativa
    doc.add_heading('2. DA JUSTIFICATIVA', 1)
    doc.add_paragraph(processo.etp_justificativa_contratacao or processo.justificativa)
    
    # 3. Especificações Técnicas
    doc.add_heading('3. DAS ESPECIFICAÇÕES TÉCNICAS', 1)
    doc.add_paragraph(processo.etp_requisitos_tecnicos_detalhe or 'Conforme especificações detalhadas no anexo.')
    
    # 4. Quantitativos
    doc.add_heading('4. DOS QUANTITATIVOS', 1)
    doc.add_paragraph(processo.etp_texto_estimativa_quantidades)
    
    # 5. Valor Estimado
    doc.add_heading('5. DO VALOR ESTIMADO', 1)
    if processo.valor_estimado:
        doc.add_paragraph(f'O valor estimado para esta contratação é de R$ {processo.valor_estimado:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'))
    
    # 6. Prazo e Local de Entrega
    doc.add_heading('6. DO PRAZO E LOCAL DE ENTREGA', 1)
    doc.add_paragraph('Os produtos/serviços deverão ser entregues conforme cronograma estabelecido no edital.')
    
    # 7. Condições de Pagamento
    doc.add_heading('7. DAS CONDIÇÕES DE PAGAMENTO', 1)
    doc.add_paragraph('O pagamento será efetuado em até 30 (trinta) dias após a entrega e aceitação dos produtos/serviços.')
    
    # 8. Obrigações da Contratada
    doc.add_heading('8. DAS OBRIGAÇÕES DA CONTRATADA', 1)
    doc.add_paragraph('A contratada deverá:')
    doc.add_paragraph('a) Fornecer os produtos/serviços conforme especificações;', style='List Bullet')
    doc.add_paragraph('b) Responsabilizar-se por todos os encargos;', style='List Bullet')
    doc.add_paragraph('c) Manter durante toda a execução as condições de habilitação.', style='List Bullet')
    
    # 9. Obrigações da Contratante
    doc.add_heading('9. DAS OBRIGAÇÕES DA CONTRATANTE', 1)
    doc.add_paragraph('A contratante deverá:')
    doc.add_paragraph('a) Efetuar o pagamento nas condições estabelecidas;', style='List Bullet')
    doc.add_paragraph('b) Fiscalizar a execução do contrato;', style='List Bullet')
    doc.add_paragraph('c) Notificar a contratada sobre irregularidades.', style='List Bullet')
    
    # 10. Fiscalização
    doc.add_heading('10. DA FISCALIZAÇÃO', 1)
    if processo.etp_fiscal_tecnico:
        doc.add_paragraph(f'Fiscal Técnico: {processo.etp_fiscal_tecnico}')
    if processo.etp_fiscal_administrativo:
        doc.add_paragraph(f'Fiscal Administrativo: {processo.etp_fiscal_administrativo}')
    if processo.etp_gestor_contrato:
        doc.add_paragraph(f'Gestor do Contrato: {processo.etp_gestor_contrato}')
    
    # 11. Vigência
    doc.add_heading('11. DA VIGÊNCIA', 1)
    if processo.vigencia_meses:
        doc.add_paragraph(f'O contrato terá vigência de {processo.vigencia_meses} meses, contados da data de sua assinatura.')
    
    # 12. Dotação Orçamentária
    doc.add_heading('12. DA DOTAÇÃO ORÇAMENTÁRIA', 1)
    if processo.etp_dotacao_programa_trabalho:
        doc.add_paragraph(f'Programa de Trabalho: {processo.etp_dotacao_programa_trabalho}')
    if processo.etp_dotacao_natureza_despesa:
        doc.add_paragraph(f'Natureza da Despesa: {processo.etp_dotacao_natureza_despesa}')
    if processo.etp_dotacao_fonte_recursos:
        doc.add_paragraph(f'Fonte de Recursos: {processo.etp_dotacao_fonte_recursos}')
    
    # Assinatura
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('Responsável Técnico')
    
    # Salvar documento
    filename = f'TR_{processo.numero_processo.replace("/", "-")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    filepath = os.path.join(settings.MEDIA_ROOT, 'documentos', filename)
    
    # Criar diretório se não existir
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc.save(filepath)
    
    # Salvar registro no banco
    documento = Documento.objects.create(
        processo=processo,
        tipo='TR',
        arquivo=f'documentos/{filename}',
        gerado_por=request.user.username
    )
    
    messages.success(request, f'Termo de Referência gerado com sucesso!')
    
    # Retornar arquivo para download
    return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=filename)

@login_required
def download_documento(request, documento_id):
    """Download de documento já gerado"""
    documento = get_object_or_404(Documento, pk=documento_id)
    filepath = os.path.join(settings.MEDIA_ROOT, str(documento.arquivo))
    
    if os.path.exists(filepath):
        return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=os.path.basename(filepath))
    else:
        messages.error(request, 'Arquivo não encontrado.')
        return redirect('documentos_list')

# ==== (FIM) CÓDIGO MESCLADO DO 'views 1.txt' ====


# --- VIEW DE EXPORTAÇÃO (Andamento PDF) ---
@login_required
def exportar_andamento_pdf(request, pk):
    processo = get_object_or_404(Processo.objects.select_related('orgao_responsavel'), pk=pk)
    historico_qs = HistoricoProcesso.objects.filter(processo=processo)
    etapas_concluidas_dict = {item.etapa: item.data_conclusao for item in historico_qs}
    todas_etapas = HistoricoProcesso.ETAPAS_CHOICES
    fase_interna_etapas = [{'key': key, 'name': name, 'date': etapas_concluidas_dict.get(key)} for key, name in todas_etapas[:12]]
    fase_externa_etapas = [{'key': key, 'name': name, 'date': etapas_concluidas_dict.get(key)} for key, name in todas_etapas[12:]]
    context = {
        'processo': processo, 
        'fase_interna_etapas': fase_interna_etapas, 
        'fase_externa_etapas': fase_externa_etapas, 
        'data_hoje': date.today().strftime("%d/%m/%Y")
    }
    html_string = render_to_string('core/processo_andamento_pdf.html', context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="andamento_processo_{processo.numero_processo}.pdf"'
    HTML(string=html_string).write_pdf(response)
    return response

# --- VIEW DA IA ---
@csrf_exempt
def generate_etp_justificativa_ia(request):
    if request.method != 'POST': 
        return JsonResponse({'error': 'Este endpoint aceita apenas requisições POST.'}, status=405)
    try:
        data = json.loads(request.body)
        necessidade = data.get('necessidade', '')
        requisitos = data.get('requisitos', '')
        levantamento = data.get('levantamento', '')
        genai.configure(api_key=settings.GOOGLE_API_KEY)
        model = genai.GenerativeModel('gemini-pro')
        prompt = f"""Você é um especialista em licitações e contratos públicos no Brasil, redigindo um Estudo Técnico Preliminar (ETP). Sua tarefa é escrever um parágrafo conciso e formal para o campo "Justificativa do Objeto da Contratação". Use as seguintes informações como base estrita para sua resposta: 1. **Necessidade da Contratação:** "{necessidade}" 2. **Requisitos da Contratação:** "{requisitos}" 3. **Levantamento de Mercado e Solução Escolhida:** "{levantamento}" Combine essas informações em um texto coeso que justifique claramente por que a contratação do objeto é necessária e a melhor solução para a administração pública. Não invente informações."""
        response = model.generate_content(prompt)
        return JsonResponse({'justificativa': response.text})
    except Exception as e:
        return JsonResponse({'error': f'Ocorreu um erro ao gerar o texto: {str(e)}'}, status=500)

# --- VIEW DO CHECKLIST ---
@csrf_exempt
@login_required
def salvar_etapa_historico(request):
    if request.method != 'POST': 
        return JsonResponse({'error': 'Método não permitido'}, status=405)
    try:
        data = json.loads(request.body)
        processo_id = data.get('processo_id')
        etapa_key = data.get('etapa')
        processo = get_object_or_404(Processo, pk=processo_id)
        historico, created = HistoricoProcesso.objects.get_or_create(processo=processo, etapa=etapa_key)
        if created:
            if etapa_key == 'PUBLICACAO_AVISO': 
                processo.status = 'PUBLICADO'
                processo.save()
            elif etapa_key == 'HOMOLOGACAO': 
                processo.status = 'HOMOLOGADO'
                processo.save()
            return JsonResponse({'success': True, 'message': 'Etapa salva com sucesso.'})
        else:
            return JsonResponse({'success': False, 'message': 'Etapa já havia sido salva.'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)

# ==============================================================================
# VIEW DE RELATÓRIOS
# ==============================================================================
@login_required
def relatorio_processos(request):
    resultados = Processo.objects.select_related('orgao_responsavel', 'secretaria_responsavel').order_by('-data_abertura')
    orgao_id = request.GET.get('orgao')
    secretaria_id = request.GET.get('secretaria')
    modalidade = request.GET.get('modalidade')
    status = request.GET.get('status')
    data_inicio = request.GET.get('data_inicio')
    data_fim = request.GET.get('data_fim')
    
    if orgao_id: 
        resultados = resultados.filter(orgao_responsavel__pk=orgao_id)
    if secretaria_id: 
        resultados = resultados.filter(secretaria_responsavel__pk=secretaria_id)
    if modalidade: 
        resultados = resultados.filter(modalidade=modalidade)
    if status: 
        resultados = resultados.filter(status=status)
    if data_inicio: 
        resultados = resultados.filter(data_abertura__gte=data_inicio)
    if data_fim: 
        resultados = resultados.filter(data_abertura__lte=data_fim)
        
    orgaos = Orgao.objects.all().order_by('nome')
    secretarias = Secretaria.objects.all().order_by('nome')
    modalidades_choices = Processo.MODALIDADE_CHOICES
    status_choices = Processo.STATUS_CHOICES
    
    context = {
        'orgaos': orgaos, 
        'secretarias': secretarias, 
        'modalidades': modalidades_choices, 
        'status_list': status_choices, 
        'resultados': resultados, 
        'filtros_aplicados': request.GET
    }
    return render(request, 'core/relatorio_processos.html', context)

# --- VIEWS DE RELATÓRIOS FALTANTES (STUBS) ---
@login_required
def relatorio_fases_excel(request):
    # Lógica para gerar Excel aqui
    messages.info(request, "Função 'Relatório Fases Excel' ainda não implementada.")
    return redirect('relatorio_processos')

@login_required
def relatorio_fases_pdf(request):
    # Lógica para gerar PDF aqui
    messages.info(request, "Função 'Relatório Fases PDF' ainda não implementada.")
    return redirect('relatorio_processos')
# --- FIM DOS STUBS DE RELATÓRIOS ---


@login_required
def exportar_processos_csv(request):
    resultados = Processo.objects.select_related('orgao_responsavel', 'secretaria_responsavel').prefetch_related('historico').order_by('-data_abertura')
    # ... (lógica de filtro repetida, idealmente seria refatorada) ...
    orgao_id = request.GET.get('orgao'); secretaria_id = request.GET.get('secretaria'); modalidade = request.GET.get('modalidade'); status = request.GET.get('status'); data_inicio = request.GET.get('data_inicio'); data_fim = request.GET.get('data_fim')
    if orgao_id: resultados = resultados.filter(orgao_responsavel__pk=orgao_id)
    if secretaria_id: resultados = resultados.filter(secretaria_responsavel__pk=secretaria_id)
    if modalidade: resultados = resultados.filter(modalidade=modalidade)
    if status: resultados = resultados.filter(status=status)
    if data_inicio: resultados = resultados.filter(data_abertura__gte=data_inicio)
    if data_fim: resultados = resultados.filter(data_abertura__lte=data_fim)
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="relatorio_processos_{date.today().strftime("%Y-%m-%d")}.csv"'
    response.write(u'\ufeff'.encode('utf8'))
    writer = csv.writer(response, delimiter=';')
    writer.writerow(['Nº Processo', 'Objeto', 'Órgão', 'Status', 'Data de Abertura', 'Valor Estimado', 'Etapas Concluídas (Andamento)'])
    
    for processo in resultados:
        etapas_concluidas = [item.get_etapa_display() for item in processo.historico.all()]
        andamento = " | ".join(etapas_concluidas) if etapas_concluidas else "Nenhuma etapa iniciada"
        writer.writerow([
            processo.numero_processo, 
            processo.objeto, 
            processo.orgao_responsavel.nome, 
            processo.get_status_display(), 
            processo.data_abertura.strftime("%d/%m/%Y"), 
            str(processo.valor_estimado).replace('.', ',') if processo.valor_estimado else '0,00', 
            andamento
        ])
    return response

@login_required
def exportar_processos_pdf(request):
    resultados = Processo.objects.select_related('orgao_responsavel', 'secretaria_responsavel').prefetch_related('historico').order_by('-data_abertura')
    # ... (lógica de filtro repetida) ...
    orgao_id = request.GET.get('orgao'); secretaria_id = request.GET.get('secretaria'); modalidade = request.GET.get('modalidade'); status = request.GET.get('status'); data_inicio = request.GET.get('data_inicio'); data_fim = request.GET.get('data_fim')
    if orgao_id: resultados = resultados.filter(orgao_responsavel__pk=orgao_id)
    if secretaria_id: resultados = resultados.filter(secretaria_responsavel__pk=secretaria_id)
    if modalidade: resultados = resultados.filter(modalidade=modalidade)
    if status: resultados = resultados.filter(status=status)
    if data_inicio: resultados = resultados.filter(data_abertura__gte=data_inicio)
    if data_fim: resultados = resultados.filter(data_abertura__lte=dataim)
    
    context = {
        'resultados': resultados, 
        'data_hoje': date.today().strftime("%d/%m/%Y"), 
        'filtros_aplicados': request.GET
    }
    html_string = render_to_string('core/relatorio_pdf_template.html', context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="relatorio_processos_{date.today().strftime("%Y-%m-%d")}.pdf"'
    HTML(string=html_string).write_pdf(response)
    return response

# ==============================================================================
# VIEWS DE ÓRGÃOS, FORNECEDORES E RESPONSÁVEIS (CRUDs)
# ==============================================================================
@login_required
def orgao_list(request):
    orgaos = Orgao.objects.all().order_by('nome')
    return render(request, 'core/orgao_list.html', {'orgaos': orgaos})

@login_required
def orgao_create(request):
    if request.method == 'POST': 
        form = OrgaoForm(request.POST, request.FILES)
    else: 
        form = OrgaoForm()
    if 'form' in locals() and form.is_valid(): 
        form.save()
        return redirect('orgao_list')
    return render(request, 'core/orgao_form.html', {'form': form, 'titulo': 'Adicionar Novo Órgão Público'})

@login_required
def orgao_update(request, pk):
    orgao = get_object_or_404(Orgao, pk=pk)
    if request.method == 'POST': 
        form = OrgaoForm(request.POST, request.FILES, instance=orgao)
    else: 
        form = OrgaoForm(instance=orgao)
    if 'form' in locals() and form.is_valid(): 
        form.save()
        return redirect('orgao_list')
    return render(request, 'core/orgao_form.html', {'form': form, 'titulo': 'Editar Órgão Público'})

@login_required
def orgao_delete(request, pk):
    orgao = get_object_or_404(Orgao, pk=pk)
    if request.method == 'POST': 
        orgao.delete()
        return redirect('orgao_list')
    return render(request, 'core/orgao_confirm_delete.html', {'orgao': orgao})

@login_required
def fornecedor_list(request):
    fornecedores = Fornecedor.objects.all().order_by('razao_social')
    return render(request, 'core/fornecedor_list.html', {'fornecedores': fornecedores})

@login_required
def fornecedor_create(request):
    if request.method == 'POST': 
        form = FornecedorForm(request.POST)
    else: 
        form = FornecedorForm()
    if 'form' in locals() and form.is_valid(): 
        form.save()
        return redirect('fornecedor_list')
    return render(request, 'core/fornecedor_form.html', {'form': form, 'titulo': 'Adicionar Novo Fornecedor'})

@login_required
def fornecedor_update(request, pk):
    fornecedor = get_object_or_404(Fornecedor, pk=pk)
    if request.method == 'POST': 
        form = FornecedorForm(request.POST, instance=fornecedor)
    else: 
        form = FornecedorForm(instance=fornecedor)
    if 'form' in locals() and form.is_valid(): 
        form.save()
        return redirect('fornecedor_list')
    return render(request, 'core/fornecedor_form.html', {'form': form, 'titulo': 'Editar Fornecedor'})

@login_required
def fornecedor_delete(request, pk):
    fornecedor = get_object_or_404(Fornecedor, pk=pk)
    if request.method == 'POST': 
        fornecedor.delete()
        return redirect('fornecedor_list')
    return render(request, 'core/fornecedor_confirm_delete.html', {'fornecedor': fornecedor})

@login_required
def responsavel_list(request):
    object_list = Responsavel.objects.select_related('secretaria').order_by('nome')
    return render(request, 'core/responsavel_list.html', {'object_list': object_list})

@login_required
def responsavel_create(request):
    if request.method == 'POST': 
        form = ResponsavelForm(request.POST)
    else: 
        form = ResponsavelForm()
    if 'form' in locals() and form.is_valid(): 
        form.save()
        return redirect('responsavel_list')
    return render(request, 'core/responsavel_form.html', {'form': form, 'titulo': 'Adicionar Novo Responsável'})

@login_required
def responsavel_update(request, pk):
    obj = get_object_or_404(Responsavel, pk=pk)
    if request.method == 'POST': 
        form = ResponsavelForm(request.POST, instance=obj)
    else: 
        form = ResponsavelForm(instance=obj)
    if 'form' in locals() and form.is_valid(): 
        form.save()
        return redirect('responsavel_list')
    return render(request, 'core/responsavel_form.html', {'form': form, 'titulo': 'Editar Responsável'})

@login_required
def responsavel_delete(request, pk):
    obj = get_object_or_404(Responsavel, pk=pk)
    if request.method == 'POST': 
        obj.delete()
        return redirect('responsavel_list')
    return render(request, 'core/responsavel_confirm_delete.html', {'object': obj})